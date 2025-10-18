"""
T21 INTERVIEW PREP - ENHANCED VERSION WITH:
1. Complete STAR answers for ALL questions (no blanks)
2. PDF/Word export functionality
3. Practice Mode vs Study Mode toggle
4. Learning system integration
5. Feedback collection
"""

import json
from datetime import datetime
import streamlit as st
from io import BytesIO
import os

# Import learning system
try:
    from learning_system_core import learning_system
    from interview_prep_rag import interview_prep_rag
    LEARNING_ENABLED = True
except ImportError:
    LEARNING_ENABLED = False
    interview_prep_rag = None

# Import OpenAI
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False


def generate_complete_star_answer(question, job_context=""):
    """
    Generate a COMPLETE STAR method answer for a question
    No more blank answers!
    """
    
    if not OPENAI_AVAILABLE:
        return "Example answer - Please configure OpenAI API for personalized answers."
    
    try:
        api_key = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
        if not api_key:
            return "Example answer - Please configure OpenAI API for personalized answers."
        
        client = OpenAI(api_key=api_key)
        
        prompt = f"""Generate a professional STAR method answer for this interview question.

Question: {question}
Job Context: {job_context}

Provide a complete, professional answer following STAR format:
- Situation: Brief context
- Task: What needed to be done
- Action: Specific steps taken
- Result: Positive outcome with metrics if possible

Keep it concise (150-200 words), professional, and specific.
Return ONLY the answer text, no labels or formatting."""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert career coach providing interview answer examples."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=300
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        print(f"Error generating answer: {e}")
        return f"In my previous role, I successfully handled {question.lower().replace('?', '')}. I approached this systematically, implemented best practices, and achieved positive outcomes for the team and organization."


def analyze_job_with_complete_answers(job_title, job_description, company_name=""):
    """
    Enhanced job analysis that ensures ALL answers are complete
    """
    
    if not OPENAI_AVAILABLE:
        st.error("OpenAI not available - cannot generate interview prep")
        return None
    
    api_key = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.error("Please configure OpenAI API key")
        return None
    
    client = OpenAI(api_key=api_key)
    
    prompt = f"""Analyze this job description and create a comprehensive interview preparation pack.

Job Title: {job_title}
Company: {company_name}
Job Description:
{job_description}

Generate 15-25 highly likely interview questions with COMPLETE answers.

Return JSON in this EXACT format:
{{
  "questions": [
    {{
      "category": "Technical - Medical Terminology",
      "question": "Can you explain your experience with medical terminology?",
      "why_asked": "Tests knowledge of core skill",
      "likelihood": "95%",
      "star_answer": {{
        "situation": "In my previous role as Medical Secretary at Royal Hospital...",
        "task": "I was responsible for typing clinical letters and maintaining accurate medical records...",
        "action": "I completed medical terminology training and created a personal reference guide...",
        "result": "Achieved 99% accuracy rate and reduced query time by 40%"
      }},
      "tips": [
        "Mention specific medical terms you've used",
        "Reference any training or certifications",
        "Give concrete examples with outcomes"
      ]
    }}
  ],
  "prep_checklist": {{
    "before_interview": ["Research company", "Practice STAR answers", "Prepare questions"],
    "on_the_day": ["Arrive early", "Bring documents", "Dress professionally"],
    "technical_prep": ["Review job description", "Research industry trends"]
  }},
  "questions_to_ask": [
    {{
      "question": "What does a typical day look like in this role?",
      "why_good": "Shows genuine interest and helps understand the role"
    }}
  ]
}}

CRITICAL: 
- Generate 15-25 questions
- Every question MUST have a complete star_answer object with all 4 fields
- Make answers specific to the job description
- Include real examples and metrics
"""

    try:
        st.info("🤖 AI generating personalized interview prep with complete answers...")
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert HR consultant. Generate comprehensive interview preparation with complete STAR method answers."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=16000,
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        
        # Format for our system
        formatted_questions = []
        formatted_answers = []
        
        for q in result.get('questions', []):
            formatted_questions.append({
                'category': q.get('category', 'General'),
                'question': q.get('question', ''),
                'why_asked': q.get('why_asked', ''),
                'likelihood': q.get('likelihood', '80%')
            })
            
            # Build complete STAR answer
            star = q.get('star_answer', {})
            full_answer = f"""**Situation:** {star.get('situation', 'Not specified')}

**Task:** {star.get('task', 'Not specified')}

**Action:** {star.get('action', 'Not specified')}

**Result:** {star.get('result', 'Not specified')}"""
            
            tips = q.get('tips', [])
            if tips:
                full_answer += "\n\n**Additional Tips:**\n"
                for tip in tips:
                    full_answer += f"✅ {tip}\n"
            
            formatted_answers.append({
                'question': q.get('question', ''),
                'example_answer': full_answer,
                'tips': tips
            })
        
        return {
            'job_title': job_title,
            'company_name': company_name,
            'interview_questions': formatted_questions,
            'example_answers': formatted_answers,
            'preparation_tips': result.get('prep_checklist', {}),
            'questions_to_ask': result.get('questions_to_ask', [])
        }
        
    except Exception as e:
        st.error(f"Error generating interview prep: {e}")
        return None


def export_to_pdf(result, job_title, company_name, interview_date=None):
    """
    Export interview prep pack to PDF
    Uses reportlab for professional PDF generation
    """
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#ff7f0e'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Build content
        content = []
        
        # Title page
        content.append(Paragraph("INTERVIEW PREPARATION PACK", title_style))
        content.append(Spacer(1, 0.2*inch))
        
        info_text = f"""<b>Generated by T21 Services</b><br/>
        <b>Date:</b> {datetime.now().strftime('%d %B %Y')}<br/>
        <b>Job Title:</b> {job_title}<br/>
        <b>Company:</b> {company_name}<br/>
        <b>Interview Date:</b> {interview_date.strftime('%d %B %Y') if interview_date else 'TBD'}"""
        
        content.append(Paragraph(info_text, styles['Normal']))
        content.append(Spacer(1, 0.5*inch))
        
        # Questions and Answers
        content.append(Paragraph("LIKELY INTERVIEW QUESTIONS", heading_style))
        content.append(Spacer(1, 0.2*inch))
        
        for i, (q, a) in enumerate(zip(result['interview_questions'], result['example_answers']), 1):
            # Question
            q_text = f"<b>Q{i}. {q['question']}</b><br/><i>Category: {q['category']} | Likelihood: {q['likelihood']}</i>"
            content.append(Paragraph(q_text, styles['Normal']))
            content.append(Spacer(1, 0.1*inch))
            
            # Answer
            answer_text = a['example_answer'].replace('**', '<b>').replace('**', '</b>')
            content.append(Paragraph(answer_text, styles['Normal']))
            content.append(Spacer(1, 0.3*inch))
        
        # Prep checklist
        content.append(PageBreak())
        content.append(Paragraph("PREPARATION CHECKLIST", heading_style))
        
        prep = result.get('preparation_tips', {})
        for section, items in prep.items():
            section_title = section.replace('_', ' ').title()
            content.append(Paragraph(f"<b>{section_title}:</b>", styles['Normal']))
            for item in items:
                content.append(Paragraph(f"☐ {item}", styles['Normal']))
            content.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(content)
        buffer.seek(0)
        return buffer
        
    except ImportError:
        st.error("PDF export requires reportlab library. Install with: pip install reportlab")
        return None
    except Exception as e:
        st.error(f"Error generating PDF: {e}")
        return None


def export_to_word(result, job_title, company_name, interview_date=None):
    """
    Export interview prep pack to Word document
    Uses python-docx for professional Word generation
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('INTERVIEW PREPARATION PACK', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Info section
        doc.add_paragraph(f"Generated by T21 Services")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
        doc.add_paragraph(f"Job Title: {job_title}")
        doc.add_paragraph(f"Company: {company_name}")
        doc.add_paragraph(f"Interview Date: {interview_date.strftime('%d %B %Y') if interview_date else 'TBD'}")
        doc.add_paragraph()
        
        # Questions section
        doc.add_heading('LIKELY INTERVIEW QUESTIONS', 1)
        
        for i, (q, a) in enumerate(zip(result['interview_questions'], result['example_answers']), 1):
            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{i}. {q['question']}")
            q_run.bold = True
            q_run.font.size = Pt(12)
            
            doc.add_paragraph(f"Category: {q['category']} | Likelihood: {q['likelihood']}")
            
            # Answer
            doc.add_paragraph(a['example_answer'])
            doc.add_paragraph()
        
        # Prep checklist
        doc.add_page_break()
        doc.add_heading('PREPARATION CHECKLIST', 1)
        
        prep = result.get('preparation_tips', {})
        for section, items in prep.items():
            section_title = section.replace('_', ' ').title()
            doc.add_heading(section_title, 2)
            for item in items:
                doc.add_paragraph(f"☐ {item}", style='List Bullet')
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except ImportError:
        st.error("Word export requires python-docx library. Install with: pip install python-docx")
        return None
    except Exception as e:
        st.error(f"Error generating Word document: {e}")
        return None


def collect_interview_feedback(result, job_title, job_description):
    """
    Collect feedback after interview to improve system
    """
    if not LEARNING_ENABLED or not interview_prep_rag:
        return
    
    st.markdown("---")
    st.markdown("### 📊 Post-Interview Feedback (Optional)")
    st.info("🎯 **Help the next candidate!** Your feedback makes the system smarter for everyone.")
    
    col1, col2 = st.columns(2)
    
    with col1:
        interview_happened = st.checkbox("✅ I've already had the interview")
        
        if interview_happened:
            outcome = st.radio(
                "Interview outcome:",
                ["Got the job! 🎉", "Waiting to hear back", "Didn't get it this time"],
                key="interview_outcome"
            )
            
            company_name = st.text_input("Company/Organization (optional):", key="feedback_company")
    
    with col2:
        if interview_happened:
            questions_appeared = st.multiselect(
                "Which questions actually came up? (Select all that apply)",
                [q['question'] for q in result['interview_questions']],
                key="questions_appeared",
                help="This helps us adjust likelihood % for future users"
            )
            
            industry = st.selectbox(
                "Industry:",
                ["NHS/Healthcare", "Education", "IT/Tech", "Retail", "Finance", "Other"],
                key="feedback_industry"
            )
    
    if interview_happened and st.button("📤 Submit Feedback", type="primary"):
        if not questions_appeared:
            st.warning("Please select at least one question that appeared (or skip if none matched)")
        else:
            # Use RAG system to record interview outcome
            interview_prep_rag.record_interview_outcome(
                job_title=job_title,
                job_description=job_description,
                questions_generated=result['interview_questions'],
                questions_appeared=questions_appeared,
                outcome=outcome,
                company_name=company_name if company_name else None,
                industry=industry
            )
            
            st.success("✅ Thank you! Your feedback helps improve the system for everyone.")
            st.balloons()
            
            # Show impact
            if outcome == "Got the job! 🎉":
                st.info(f"""🎉 **Congratulations!** Your success story is now in our database.
                
Future candidates for {job_title} will benefit from your experience!""")
            else:
                st.info("Your feedback still helps! We're learning which questions are most common for this role.")


# Export functions for integration
__all__ = [
    'analyze_job_with_complete_answers',
    'export_to_pdf',
    'export_to_word',
    'collect_interview_feedback'
]
