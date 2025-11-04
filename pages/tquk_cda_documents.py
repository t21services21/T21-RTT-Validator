"""
TQUK CDA SUBMISSION DOCUMENTS
Download page for all CDA submission materials
"""

import streamlit as st
import os
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

st.set_page_config(page_title="TQUK CDA Documents", page_icon="📋", layout="wide")

# Conversion functions
def markdown_to_word(markdown_text, title):
    """Convert markdown to Word document"""
    doc = Document()
    
    # Add title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Centre: T21 Services UK (#36257481088)")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("")
    
    # Process markdown content
    lines = markdown_text.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 4)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif line.strip().startswith('✓ ') or line.strip().startswith('✅ '):
            doc.add_paragraph(line.strip(), style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def markdown_to_pdf(markdown_text, title):
    """Convert markdown to PDF document"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 0, 139),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=RGBColor(0, 0, 139),
        spaceAfter=12,
        spaceBefore=12
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=RGBColor(0, 0, 139),
        spaceAfter=10,
        spaceBefore=10
    )
    
    # Add title
    elements.append(Paragraph(title, title_style))
    elements.append(Spacer(1, 12))
    
    # Add metadata
    elements.append(Paragraph(f"<b>Centre:</b> T21 Services UK (#36257481088)", styles['Normal']))
    elements.append(Paragraph(f"<b>Generated:</b> {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Process markdown content
    lines = markdown_text.split('\n')
    for line in lines:
        if line.startswith('# '):
            elements.append(Paragraph(line[2:], heading1_style))
        elif line.startswith('## '):
            elements.append(Paragraph(line[3:], heading2_style))
        elif line.startswith('### '):
            elements.append(Paragraph(line[4:], styles['Heading3']))
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            elements.append(Paragraph(f"• {line.strip()[2:]}", styles['Normal']))
        elif line.strip().startswith('✓ ') or line.strip().startswith('✅ '):
            elements.append(Paragraph(line.strip(), styles['Normal']))
        elif line.strip() == '---':
            elements.append(Spacer(1, 12))
        elif line.strip():
            # Clean up markdown formatting
            clean_line = line.replace('**', '').replace('`', '')
            elements.append(Paragraph(clean_line, styles['Normal']))
            elements.append(Spacer(1, 6))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer

# Header
st.title("📋 TQUK CDA Submission Documents")
st.markdown("**Centre:** T21 Services UK (#36257481088)")
st.markdown(f"**Generated:** {datetime.now().strftime('%B %d, %Y')}")

st.divider()

# Introduction
st.info("""
**✅ UPDATED: Level 3 Diploma in Adult Care - Complete CDA Submission**

This page provides access to ALL Centre-Devised Assessment (CDA) submission documents 
for TQUK Level 3 Diploma in Adult Care (RQF) – 610/0103/6

**NEW:** All 7 mandatory unit assessment plans with GLH breakdown and assessment methods!
""")

# Document list
st.header("📄 Available Documents")

col1, col2 = st.columns(2)

with col1:
    st.subheader("1️⃣ Master Mapping Document")
    st.markdown("""
    **Size:** 6,500+ words  
    **Content:**
    - Complete mapping of all 10 qualifications
    - Learning outcomes coverage
    - Assessment criteria alignment
    - Platform overview
    - Evidence requirements
    """)
    
    # Check if file exists
    file_path_1 = "TQUK_CDA_MASTER_MAPPING_DOCUMENT.md"
    if os.path.exists(file_path_1):
        with open(file_path_1, 'r', encoding='utf-8') as f:
            content_1 = f.read()
        
        # PDF Download
        pdf_buffer_1 = markdown_to_pdf(content_1, "TQUK CDA Master Mapping Document")
        st.download_button(
            label="📄 Download as PDF",
            data=pdf_buffer_1,
            file_name="TQUK_CDA_Master_Mapping_Document.pdf",
            mime="application/pdf",
            use_container_width=True
        )
        
        # Word Download
        word_buffer_1 = markdown_to_word(content_1, "TQUK CDA Master Mapping Document")
        st.download_button(
            label="📝 Download as Word",
            data=word_buffer_1,
            file_name="TQUK_CDA_Master_Mapping_Document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.error("File not found")

with col2:
    st.subheader("2️⃣ Assessment Strategy")
    st.markdown("""
    **Size:** 4,500+ words  
    **Content:**
    - Assessment principles (VRASAC)
    - Assessment methods detailed
    - Quality assurance processes
    - Staff roles and responsibilities
    - Continuous improvement
    """)
    
    file_path_2 = "TQUK_CDA_ASSESSMENT_STRATEGY.md"
    if os.path.exists(file_path_2):
        with open(file_path_2, 'r', encoding='utf-8') as f:
            content_2 = f.read()
        
        # PDF Download
        pdf_buffer_2 = markdown_to_pdf(content_2, "TQUK CDA Assessment Strategy")
        st.download_button(
            label="📄 Download as PDF",
            data=pdf_buffer_2,
            file_name="TQUK_CDA_Assessment_Strategy.pdf",
            mime="application/pdf",
            use_container_width=True
        )
        
        # Word Download
        word_buffer_2 = markdown_to_word(content_2, "TQUK CDA Assessment Strategy")
        st.download_button(
            label="📝 Download as Word",
            data=word_buffer_2,
            file_name="TQUK_CDA_Assessment_Strategy.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.error("File not found")

st.divider()

col3, col4 = st.columns(2)

with col3:
    st.subheader("3️⃣ Form Completion Guide")
    st.markdown("""
    **Size:** 3,000+ words  
    **Content:**
    - Step-by-step form filling instructions
    - Exact text to copy/paste
    - Email template
    - Submission checklist
    - Timeline expectations
    """)
    
    file_path_3 = "TQUK_CDA_FORM_COMPLETION_GUIDE.md"
    if os.path.exists(file_path_3):
        with open(file_path_3, 'r', encoding='utf-8') as f:
            content_3 = f.read()
        
        # PDF Download
        pdf_buffer_3 = markdown_to_pdf(content_3, "TQUK CDA Form Completion Guide")
        st.download_button(
            label="📄 Download as PDF",
            data=pdf_buffer_3,
            file_name="TQUK_CDA_Form_Completion_Guide.pdf",
            mime="application/pdf",
            use_container_width=True
        )
        
        # Word Download
        word_buffer_3 = markdown_to_word(content_3, "TQUK CDA Form Completion Guide")
        st.download_button(
            label="📝 Download as Word",
            data=word_buffer_3,
            file_name="TQUK_CDA_Form_Completion_Guide.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.error("File not found")

with col4:
    st.subheader("4️⃣ Complete Package Summary")
    st.markdown("""
    **Size:** 2,500+ words  
    **Content:**
    - Overview of everything
    - Your action steps
    - Timeline
    - Success metrics
    - Revenue potential
    """)
    
    file_path_4 = "TQUK_CDA_SUBMISSION_COMPLETE_PACKAGE.md"
    if os.path.exists(file_path_4):
        with open(file_path_4, 'r', encoding='utf-8') as f:
            content_4 = f.read()
        
        # PDF Download
        pdf_buffer_4 = markdown_to_pdf(content_4, "TQUK CDA Complete Package Summary")
        st.download_button(
            label="📄 Download as PDF",
            data=pdf_buffer_4,
            file_name="TQUK_CDA_Complete_Package_Summary.pdf",
            mime="application/pdf",
            use_container_width=True
        )
        
        # Word Download
        word_buffer_4 = markdown_to_word(content_4, "TQUK CDA Complete Package Summary")
        st.download_button(
            label="📝 Download as Word",
            data=word_buffer_4,
            file_name="TQUK_CDA_Complete_Package_Summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.error("File not found")

st.divider()

# NEW SECTION: Level 3 Adult Care Unit Assessment Plans
st.header("🎓 Level 3 Diploma in Adult Care - Unit Assessment Plans")
st.success("✅ **COMPLETE SUBMISSION:** All 7 mandatory units with comprehensive assessment plans!")

# Base path for unit files
unit_base_path = "TQUK_CDA_Complete_Submission"

# Unit details
units = [
    {
        "number": "1",
        "title": "Duty of Care",
        "file": "Unit_1_Assessment_Plan_Duty_of_Care.md",
        "glh": "20",
        "credit": "2"
    },
    {
        "number": "2",
        "title": "Equality, Diversity & Inclusion",
        "file": "Unit_2_Assessment_Plan_Equality_Diversity_Inclusion.md",
        "glh": "20",
        "credit": "2"
    },
    {
        "number": "3",
        "title": "Person-Centred Practice",
        "file": "Unit_3_Assessment_Plan_Person_Centred_Practice.md",
        "glh": "20",
        "credit": "2"
    },
    {
        "number": "4",
        "title": "Safeguarding",
        "file": "Unit_4_Assessment_Plan_Safeguarding.md",
        "glh": "20",
        "credit": "2"
    },
    {
        "number": "5",
        "title": "Communication",
        "file": "Unit_5_Assessment_Plan_Communication.md",
        "glh": "20",
        "credit": "2"
    },
    {
        "number": "6",
        "title": "Health & Safety",
        "file": "Unit_6_Assessment_Plan_Health_Safety.md",
        "glh": "20",
        "credit": "2"
    },
    {
        "number": "7",
        "title": "CPD",
        "file": "Unit_7_Assessment_Plan_CPD.md",
        "glh": "20",
        "credit": "2"
    }
]

# Display units in 2 columns
for i in range(0, len(units), 2):
    col_left, col_right = st.columns(2)
    
    # Left column
    with col_left:
        unit = units[i]
        st.subheader(f"Unit {unit['number']}: {unit['title']}")
        st.markdown(f"**GLH:** {unit['glh']} | **Credit:** {unit['credit']}")
        
        file_path = os.path.join(unit_base_path, unit['file'])
        if os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # PDF Download
            pdf_buffer = markdown_to_pdf(content, f"Unit {unit['number']}: {unit['title']}")
            st.download_button(
                label="📄 Download PDF",
                data=pdf_buffer,
                file_name=f"Unit_{unit['number']}_{unit['title'].replace(' ', '_')}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key=f"pdf_{unit['number']}"
            )
            
            # Word Download
            word_buffer = markdown_to_word(content, f"Unit {unit['number']}: {unit['title']}")
            st.download_button(
                label="📝 Download Word",
                data=word_buffer,
                file_name=f"Unit_{unit['number']}_{unit['title'].replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key=f"word_{unit['number']}"
            )
        else:
            st.error(f"File not found: {file_path}")
    
    # Right column (if exists)
    if i + 1 < len(units):
        with col_right:
            unit = units[i + 1]
            st.subheader(f"Unit {unit['number']}: {unit['title']}")
            st.markdown(f"**GLH:** {unit['glh']} | **Credit:** {unit['credit']}")
            
            file_path = os.path.join(unit_base_path, unit['file'])
            if os.path.exists(file_path):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # PDF Download
                pdf_buffer = markdown_to_pdf(content, f"Unit {unit['number']}: {unit['title']}")
                st.download_button(
                    label="📄 Download PDF",
                    data=pdf_buffer,
                    file_name=f"Unit_{unit['number']}_{unit['title'].replace(' ', '_')}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"pdf_{unit['number']}"
                )
                
                # Word Download
                word_buffer = markdown_to_word(content, f"Unit {unit['number']}: {unit['title']}")
                st.download_button(
                    label="📝 Download Word",
                    data=word_buffer,
                    file_name=f"Unit_{unit['number']}_{unit['title'].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"word_{unit['number']}"
                )
            else:
                st.error(f"File not found: {file_path}")

st.divider()

# Submission Summary and Email Response
st.header("📧 Submission Documents")

col_sum1, col_sum2 = st.columns(2)

with col_sum1:
    st.subheader("📋 CDA Submission Summary")
    st.markdown("**Complete overview of all 7 units**")
    
    summary_file = os.path.join(unit_base_path, "CDA_SUBMISSION_SUMMARY.md")
    if os.path.exists(summary_file):
        with open(summary_file, 'r', encoding='utf-8') as f:
            summary_content = f.read()
        
        # PDF Download
        pdf_summary = markdown_to_pdf(summary_content, "CDA Submission Summary")
        st.download_button(
            label="📄 Download PDF",
            data=pdf_summary,
            file_name="CDA_Submission_Summary.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="pdf_summary"
        )
        
        # Word Download
        word_summary = markdown_to_word(summary_content, "CDA Submission Summary")
        st.download_button(
            label="📝 Download Word",
            data=word_summary,
            file_name="CDA_Submission_Summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="word_summary"
        )
    else:
        st.error("Summary file not found")

with col_sum2:
    st.subheader("📧 Email Response to TQUK")
    st.markdown("**Ready-to-send email template**")
    
    email_file = os.path.join(unit_base_path, "EMAIL_RESPONSE_TO_TQUK.md")
    if os.path.exists(email_file):
        with open(email_file, 'r', encoding='utf-8') as f:
            email_content = f.read()
        
        # PDF Download
        pdf_email = markdown_to_pdf(email_content, "Email Response to TQUK")
        st.download_button(
            label="📄 Download PDF",
            data=pdf_email,
            file_name="Email_Response_to_TQUK.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="pdf_email"
        )
        
        # Word Download
        word_email = markdown_to_word(email_content, "Email Response to TQUK")
        st.download_button(
            label="📝 Download Word",
            data=word_email,
            file_name="Email_Response_to_TQUK.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="word_email"
        )
    else:
        st.error("Email file not found")

st.divider()

# Qualifications covered
st.header("📚 Qualifications Covered (10 Total)")

col5, col6, col7 = st.columns(3)

with col5:
    st.subheader("🔤 Functional Skills (4)")
    st.markdown("""
    1. English Level 1 (610/2626/6)
    2. English Level 2 (610/2625/6)
    3. Maths Level 1 (610/2623/2)
    4. Maths Level 2 (610/2624/4)
    """)

with col6:
    st.subheader("📊 Level 2 Certificates (4)")
    st.markdown("""
    5. Customer Service (603/3895/7)
    6. IT User Skills (603/5640/9)
    7. Adult Social Care (601/4040/9)
    8. Business Admin (603/2949/X)
    """)

with col7:
    st.subheader("🎓 Level 3 Qualifications (2)")
    st.markdown("""
    9. Teaching & Learning (601/2731/4)
    10. Adult Care (610/0103/6)
    """)

st.divider()

# Next steps
st.header("🚀 Next Steps")

st.success("""
**After Downloading:**

1. **Convert to PDF** - Open each .md file in Word or use an online converter
2. **Create Google Drive Folder** - Organize all documents
3. **Take Platform Screenshots** - Capture key features
4. **Fill CDA Form** - Use the Form Completion Guide
5. **Submit to TQUK** - Email with Google Drive link

**Expected Timeline:**
- Submission: Today
- TQUK Review: 5-15 working days
- Approval: December 2025
""")

# Additional resources
st.header("📖 Additional Resources")

with st.expander("📋 Qualifications List"):
    file_path_5 = "TQUK_ALL_QUALIFICATIONS_COMPLETE_LIST.md"
    if os.path.exists(file_path_5):
        with open(file_path_5, 'r', encoding='utf-8') as f:
            content_5 = f.read()
        
        st.download_button(
            label="📥 Download Complete Qualifications List",
            data=content_5,
            file_name="TQUK_All_Qualifications_List.md",
            mime="text/markdown"
        )
    else:
        st.warning("Qualifications list not found")

with st.expander("📝 Submission Guide"):
    file_path_6 = "TQUK_CDA_FORM_SUBMISSION_GUIDE.md"
    if os.path.exists(file_path_6):
        with open(file_path_6, 'r', encoding='utf-8') as f:
            content_6 = f.read()
        
        st.download_button(
            label="📥 Download Original Submission Guide",
            data=content_6,
            file_name="TQUK_CDA_Form_Submission_Guide.md",
            mime="text/markdown"
        )

st.divider()

# Contact info
st.info("""
**Questions?**

Contact TQUK:
- Email: support@tquk.org / quality@tquk.org
- Phone: 03333 58 3344

Centre Contact:
- Amb. Tosin Michael Owonifari
- Email: t.owonifari@t21services.co.uk
""")

# Footer
st.markdown("---")
st.caption("T21 Services UK | Centre #36257481088 | TQUK Approved Centre")
st.caption(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
