"""
T21 CAREER SERVICES - Professional CV Builder
AI-Powered CV Generator for ALL Career Paths

Creates ATS-optimized, professional CVs with:
- Auto-populated T21 qualifications
- Career-specific templates
- Skills matching
- Professional formatting
- PDF export
"""

from datetime import datetime
from io import BytesIO

def generate_cv_data(student_info, work_history, qualifications, skills, template='healthcare'):
    """
    Generate structured CV data from student information
    """
    
    # Ensure qualifications have description field
    enhanced_quals = []
    for qual in qualifications:
        enhanced_qual = qual.copy()
        if 'description' not in enhanced_qual:
            # Add description based on title
            title_lower = qual.get('title', '').lower()
            if 'rtt' in title_lower:
                enhanced_qual['description'] = 'RTT pathway management, clock management, NHS performance standards'
            elif 'information governance' in title_lower or 'ig' in title_lower:
                enhanced_qual['description'] = 'GDPR compliance, data protection, NHS Caldicott principles, cyber security'
            elif 'hospital administration' in title_lower:
                enhanced_qual['description'] = 'Patient administration, PAS systems, appointment booking, waiting list management'
            elif 'cancer' in title_lower:
                enhanced_qual['description'] = '62-day target, 2-week wait, cancer pathway tracking, MDT coordination'
            else:
                enhanced_qual['description'] = ''
        enhanced_quals.append(enhanced_qual)
    
    # Ensure all required fields
    cv_data = {
        'personal_info': {
            'name': student_info.get('name', ''),
            'email': student_info.get('email', ''),
            'phone': student_info.get('phone', ''),
            'location': student_info.get('location', ''),
            'linkedin': student_info.get('linkedin', ''),
            'summary': student_info.get('summary', '')
        },
        'work_history': work_history,
        'qualifications': enhanced_quals,
        'skills': skills,
        'template': template,
        'generated_date': datetime.now().strftime('%d/%m/%Y'),
        'achievements': []  # Will be populated with default achievements
    }
    
    return cv_data


def generate_professional_summary(career_path, years_experience, key_skills):
    """
    Generate professional summary based on career path
    """
    
    years_text = f"{years_experience} years" if years_experience > 0 else "experience"
    
    summaries = {
        # NHS PATHWAY & ADMIN ROLES
        'rtt_validation': f"""Detail-oriented RTT Validation Officer with {years_text} ensuring NHS compliance with 18-week targets. Expert in RTT pathway management, PAS systems, and data validation. Strong understanding of NHS performance standards, clock management, and breach prevention. Proficient in {', '.join(key_skills[:3])}. Committed to maintaining data quality and supporting patient care excellence through accurate RTT pathway tracking.""",
        
        'pathway_navigator': f"""Dedicated Patient Pathway Navigator with {years_text} coordinating patient journeys and ensuring timely access to treatment. Experienced in RTT pathway tracking, appointment scheduling, and patient communication. Skilled in {', '.join(key_skills[:3])}. Strong understanding of NHS 18-week targets and commitment to delivering exceptional patient support while managing complex pathways efficiently.""",
        
        'pathway_coordinator': f"""Organized Pathway Coordinator with {years_text} managing patient pathways and ensuring NHS performance compliance. Proficient in waiting list coordination, appointment booking, and data management using PAS systems. Skilled in {', '.join(key_skills[:3])}. Excellent communication and multitasking abilities with strong attention to detail and commitment to patient care standards.""",
        
        'cancer_tracker': f"""Meticulous Cancer Pathway Tracker with {years_text} ensuring compliance with cancer waiting time standards including 62-day and 2-week wait targets. Experienced in MDT coordination, treatment tracking, and cancer data management. Proficient in {', '.join(key_skills[:3])}. Strong understanding of NCRAS reporting and commitment to supporting cancer patients through accurate pathway monitoring.""",
        
        'waiting_list_admin': f"""Efficient Waiting List Administrator with {years_text} managing patient waiting lists and ensuring RTT compliance. Skilled in appointment scheduling, priority management, and PAS system administration. Proficient in {', '.join(key_skills[:3])}. Strong organizational skills with attention to detail and commitment to optimizing patient access to healthcare services.""",
        
        'booking_officer': f"""Professional Booking Officer with {years_text} coordinating patient appointments and managing clinic schedules. Experienced in PAS systems, Choose & Book, and e-Referrals management. Skilled in {', '.join(key_skills[:3])}. Excellent telephone manner and customer service skills with ability to manage multiple bookings efficiently while maintaining accuracy.""",
        
        'appointment_admin': f"""Reliable Appointment Administrator with {years_text} managing patient scheduling and clinic coordination. Proficient in PAS systems, diary management, and patient communication. Skilled in {', '.join(key_skills[:3])}. Strong administrative abilities with attention to detail and commitment to delivering excellent patient service in a busy healthcare environment.""",
        
        'mdt_admin': f"""Organized MDT Administrator with {years_text} coordinating multi-disciplinary team meetings and supporting cancer pathway management. Experienced in minute taking, action tracking, and patient list management. Skilled in {', '.join(key_skills[:3])}. Excellent communication and organizational abilities with understanding of medical terminology and confidentiality requirements.""",
        
        'medical_secretary': f"""Professional Medical Secretary with {years_text} providing comprehensive secretarial support to clinical teams. Experienced in clinic correspondence, audio typing, and appointment management. Proficient in {', '.join(key_skills[:3])}. Strong understanding of medical terminology and NHS systems with excellent organizational and communication skills.""",
        
        'data_officer': f"""Analytical Data Quality Officer with {years_text} ensuring accuracy and integrity of NHS information systems. Experienced in data validation, database management, and performance reporting. Skilled in {', '.join(key_skills[:3])}. Strong attention to detail with understanding of information governance and commitment to maintaining high data quality standards.""",
        
        'nhs_clerical': f"""Efficient NHS Clerical Officer with {years_text} providing administrative support in healthcare settings. Experienced in patient records management, data entry, and reception duties. Skilled in {', '.join(key_skills[:3])}. Strong organizational and communication abilities with understanding of NHS systems and commitment to maintaining confidentiality.""",
        
        # HEALTHCARE & CARE
        'healthcare_assistant': f"""Compassionate Healthcare Assistant with {years_text} providing high-quality patient care. Skilled in {', '.join(key_skills[:3])}. Committed to maintaining dignity, respect, and person-centered care. Strong understanding of safeguarding procedures and clinical observations. Seeking to contribute to a caring team in a dynamic healthcare environment.""",
        
        'care_worker': f"""Experienced Care Worker with {years_text} providing compassionate support to service users. Proficient in personal care, medication administration, and daily living support. Skilled in {', '.join(key_skills[:3])}. Strong advocate for dignity, independence, and person-centered care with excellent communication skills and commitment to safeguarding.""",
        
        # EDUCATION
        'teaching_assistant': f"""Enthusiastic Teaching Assistant with {years_text} supporting children's learning and development. Experienced in SEN/SEND support, behavior management, and differentiated learning activities. Skilled in {', '.join(key_skills[:3])}. Strong commitment to safeguarding and inclusive education with passion for helping every child reach their potential.""",
        
        # CUSTOMER SERVICE
        'customer_service': f"""Professional Customer Service Specialist with {years_text} delivering excellent customer experiences. Expert in {', '.join(key_skills[:3])}. Proven ability to handle challenging situations with empathy and professionalism. Strong communication and problem-solving skills with commitment to service excellence.""",
        
        # BUSINESS ADMIN
        'business_admin': f"""Organized Business Administrator with {years_text} supporting efficient office operations. Proficient in {', '.join(key_skills[:3])}. Strong attention to detail and excellent organizational skills with ability to manage multiple priorities and maintain accurate records.""",
        
        # TECH ROLES
        'data_analyst': f"""Results-driven Data Analyst with {years_text} transforming data into actionable insights. Proficient in {', '.join(key_skills[:3])}. Strong analytical and problem-solving abilities with expertise in data visualization and reporting. Committed to delivering data-driven solutions that support business objectives.""",
        
        'data_scientist': f"""Innovative Data Scientist with {years_text} developing machine learning solutions and predictive models. Expert in {', '.join(key_skills[:3])}. Strong statistical analysis and programming skills with ability to extract insights from complex datasets and communicate findings to stakeholders.""",
        
        'software_tester': f"""Detail-oriented Software Tester with {years_text} ensuring quality through comprehensive testing strategies. Experienced in {', '.join(key_skills[:3])}. Strong analytical skills with expertise in test automation and defect management. Committed to delivering high-quality software products.""",
        
        'business_analyst': f"""Strategic Business Analyst with {years_text} bridging business needs and technical solutions. Skilled in {', '.join(key_skills[:3])}. Strong stakeholder management and analytical abilities with expertise in requirements gathering and process improvement.""",
        
        'project_manager': f"""Dynamic Project Manager with {years_text} delivering projects on time and within budget. Experienced in {', '.join(key_skills[:3])}. Strong leadership and communication skills with proven track record of managing cross-functional teams and stakeholder expectations.""",
        
        'it_support': f"""Responsive IT Support Specialist with {years_text} providing technical assistance and troubleshooting. Proficient in {', '.join(key_skills[:3])}. Strong customer service orientation with ability to explain technical concepts clearly and resolve issues efficiently.""",
        
        'web_developer': f"""Creative Web Developer with {years_text} building responsive and user-friendly websites. Skilled in {', '.join(key_skills[:3])}. Strong frontend development abilities with understanding of UX principles and commitment to writing clean, maintainable code.""",
        
        'python_developer': f"""Proficient Python Developer with {years_text} developing scalable applications and automation solutions. Expert in {', '.join(key_skills[:3])}. Strong programming fundamentals with experience in backend development and API integration.""",
        
        # BUSINESS & PROFESSIONAL
        'digital_marketing': f"""Strategic Digital Marketing Specialist with {years_text} driving online engagement and brand growth. Experienced in {', '.join(key_skills[:3])}. Strong analytical and creative abilities with proven track record of delivering successful digital campaigns.""",
        
        'hr_officer': f"""Professional HR Officer with {years_text} supporting employee lifecycle and organizational development. Skilled in {', '.join(key_skills[:3])}. Strong understanding of employment law and HR best practices with excellent interpersonal and communication skills.""",
        
        'bookkeeper': f"""Accurate Bookkeeper with {years_text} maintaining financial records and ensuring compliance. Proficient in {', '.join(key_skills[:3])}. Strong attention to detail with excellent numerical skills and understanding of accounting principles.""",
        
        'team_leader': f"""Motivational Team Leader with {years_text} driving performance and developing talent. Experienced in {', '.join(key_skills[:3])}. Strong leadership and coaching abilities with proven track record of achieving targets and building high-performing teams."""
    }
    
    return summaries.get(career_path, f"""Motivated professional with {years_text} in {career_path.replace('_', ' ')}. Skilled in {', '.join(key_skills[:3]) if key_skills else 'various professional competencies'}. Strong work ethic, excellent communication skills, and commitment to delivering high-quality results.""")


def format_cv_html(cv_data):
    """
    Generate HTML formatted CV
    """
    
    personal = cv_data['personal_info']
    
    html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: 'Calibri', 'Arial', sans-serif;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        .header {{
            text-align: center;
            border-bottom: 3px solid #005EB8;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        .name {{
            font-size: 32px;
            font-weight: bold;
            color: #005EB8;
            margin-bottom: 10px;
        }}
        .contact {{
            font-size: 14px;
            color: #666;
        }}
        .contact a {{
            color: #005EB8;
            text-decoration: none;
        }}
        .section {{
            margin-bottom: 30px;
        }}
        .section-title {{
            font-size: 20px;
            font-weight: bold;
            color: #005EB8;
            border-bottom: 2px solid #ddd;
            padding-bottom: 5px;
            margin-bottom: 15px;
        }}
        .job {{
            margin-bottom: 20px;
        }}
        .job-title {{
            font-weight: bold;
            font-size: 16px;
        }}
        .job-company {{
            font-style: italic;
            color: #666;
        }}
        .job-dates {{
            color: #888;
            font-size: 14px;
        }}
        .job-description {{
            margin-top: 8px;
        }}
        .skills {{
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
        }}
        .skill {{
            background-color: #f0f8ff;
            border: 1px solid #005EB8;
            padding: 5px 15px;
            border-radius: 15px;
            font-size: 14px;
        }}
        .qualification {{
            margin-bottom: 15px;
        }}
        .qual-title {{
            font-weight: bold;
        }}
        .qual-institution {{
            color: #666;
            font-style: italic;
        }}
        ul {{
            margin: 10px 0;
            padding-left: 20px;
        }}
        li {{
            margin-bottom: 5px;
        }}
        .footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            font-size: 12px;
            color: #888;
        }}
    </style>
</head>
<body>
    <!-- HEADER -->
    <div class="header">
        <div class="name">{personal['name']}</div>
        <div class="contact">
            {personal['email']} | {personal['phone']} | {personal['location']}
            {' | <a href="' + personal['linkedin'] + '">LinkedIn</a>' if personal.get('linkedin') else ''}
        </div>
    </div>
    
    <!-- PROFESSIONAL SUMMARY -->
    <div class="section">
        <div class="section-title">PROFESSIONAL SUMMARY</div>
        <p>{personal['summary']}</p>
    </div>
    
    <!-- KEY SKILLS -->
    <div class="section">
        <div class="section-title">KEY SKILLS</div>
        <div class="skills">
"""
    
    for skill in cv_data['skills']:
        html += f'            <div class="skill">{skill}</div>\n'
    
    html += """        </div>
    </div>
    
    <!-- WORK EXPERIENCE -->
    <div class="section">
        <div class="section-title">WORK EXPERIENCE</div>
"""
    
    for job in cv_data['work_history']:
        html += f"""        <div class="job">
            <div class="job-title">{job['title']}</div>
            <div class="job-company">{job['company']} | <span class="job-dates">{job['dates']}</span></div>
            <div class="job-description">
                <ul>
"""
        for responsibility in job['responsibilities']:
            html += f"                    <li>{responsibility}</li>\n"
        
        html += """                </ul>
            </div>
        </div>
"""
    
    html += """    </div>
    
    <!-- EDUCATION -->
    <div class="section">
        <div class="section-title">EDUCATION</div>
"""
    
    # Add education entries (from qualifications that are degrees/diplomas)
    education_keywords = ['degree', 'diploma', 'gcse', 'a-level', 'btec', 'nvq', 'bachelor', 'master', 'phd', 'doctorate']
    education_quals = [q for q in cv_data['qualifications'] if any(keyword in q['title'].lower() for keyword in education_keywords)]
    
    if education_quals:
        for qual in education_quals:
            html += f"""        <div class="qualification">
            <div class="qual-title">{qual['title']}</div>
            <div class="qual-institution">{qual['institution']} | {qual['date']}</div>
            {f'<p>{qual["description"]}</p>' if qual.get('description') else ''}
        </div>
"""
    else:
        # If no education found, add placeholder
        html += """        <div class="qualification">
            <div class="qual-title">Secondary Education</div>
            <div class="qual-institution">GCSEs / A-Levels / Equivalent</div>
        </div>
"""
    
    html += """    </div>
    
    <!-- PROFESSIONAL QUALIFICATIONS & CERTIFICATIONS -->
    <div class="section">
        <div class="section-title">PROFESSIONAL QUALIFICATIONS & CERTIFICATIONS</div>
"""
    
    # Add professional certifications (non-education qualifications)
    cert_quals = [q for q in cv_data['qualifications'] if not any(keyword in q['title'].lower() for keyword in education_keywords)]
    
    if cert_quals:
        for qual in cert_quals:
            html += f"""        <div class="qualification">
            <div class="qual-title">{qual['title']}</div>
            <div class="qual-institution">{qual['institution']} | {qual['date']}</div>
            {f'<p>{qual["description"]}</p>' if qual.get('description') else ''}
        </div>
"""
    else:
        # Show all qualifications if no separation
        for qual in cv_data['qualifications']:
            html += f"""        <div class="qualification">
            <div class="qual-title">{qual['title']}</div>
            <div class="qual-institution">{qual['institution']} | {qual['date']}</div>
            {f'<p>{qual["description"]}</p>' if qual.get('description') else ''}
        </div>
"""
    
    html += """    </div>
    
    <!-- ACHIEVEMENTS & AWARDS -->
    <div class="section">
        <div class="section-title">KEY ACHIEVEMENTS & AWARDS</div>
        <ul>
"""
    
    # Add achievements if provided
    if cv_data.get('achievements'):
        for achievement in cv_data['achievements']:
            html += f"            <li>{achievement}</li>\n"
    else:
        # Default achievements based on qualifications
        html += """            <li>Successfully completed TQUK-Endorsed Professional Development Course in RTT & Hospital Administration</li>
            <li>Demonstrated strong commitment to professional development and continuous learning</li>
            <li>Proven track record of delivering high-quality work and meeting deadlines</li>
"""
    
    html += """        </ul>
    </div>
    
    <!-- ADDITIONAL INFORMATION -->
    <div class="section">
        <div class="section-title">ADDITIONAL INFORMATION</div>
        <ul>
            <li><strong>Right to Work:</strong> Eligible to work in the UK</li>
            <li><strong>DBS Check:</strong> Enhanced DBS check available upon request</li>
            <li><strong>References:</strong> Available upon request</li>
            <li><strong>Languages:</strong> English (Fluent)</li>
        </ul>
    </div>
    
    <!-- FOOTER -->
    <div class="footer">
        Professional CV generated by T21 Career Services | {date}
    </div>
</body>
</html>
""".format(date=cv_data['generated_date'])
    
    return html


def get_ats_keywords(career_path):
    """
    Get ATS (Applicant Tracking System) keywords for career path
    """
    
    keywords = {
        # NHS PATHWAY & ADMIN ROLES
        'rtt_validation': [
            'RTT Pathway', '18-Week Target', 'NHS Performance', 'PAS Systems',
            'Data Validation', 'Clock Management', 'Breach Prevention',
            'Patient Administration', 'Waiting List Management', 'Compliance',
            'Audit', 'Data Quality', 'Excel', 'Reporting', 'NHS Standards',
            'RTT Codes', 'First Definitive Treatment', 'Clock Stops', 'Referrals'
        ],
        'pathway_navigator': [
            'Patient Pathway', 'RTT Navigator', 'Patient Journey', 'Coordination',
            'Appointment Scheduling', 'Patient Communication', 'Pathway Tracking',
            'NHS 18 Weeks', 'PAS Systems', 'Waiting List', 'Clinical Liaison',
            'Patient Support', 'Referral Management', 'Breach Management', 'Data Entry'
        ],
        'pathway_coordinator': [
            'Pathway Coordination', 'Patient Pathways', 'RTT Management', 'Scheduling',
            'NHS Performance', 'Waiting List Coordination', 'Appointment Booking',
            'Data Management', 'Patient Administration', 'PAS', 'Excel', 'Reporting',
            'Communication Skills', 'Time Management', 'Multitasking', 'Attention to Detail'
        ],
        'cancer_tracker': [
            'Cancer Pathway', '62-Day Target', '2-Week Wait', 'Fast Track Referrals',
            'Cancer Data', 'MDT Coordination', 'Treatment Tracking', 'Data Quality',
            'NHS Performance', 'Patient Pathways', 'NCRAS', 'Somerset', 'Reporting',
            'Excel', 'Database Management', 'Confidentiality', 'GDPR Compliance'
        ],
        'waiting_list_admin': [
            'Waiting List Management', 'Patient Administration', 'RTT Compliance',
            'Appointment Scheduling', 'PAS Systems', 'Data Entry', 'Priority Management',
            'NHS Performance', 'Patient Communication', 'Booking Systems', 'Excel',
            'Time Management', 'Attention to Detail', 'Organizational Skills', 'GDPR'
        ],
        'booking_officer': [
            'Appointment Booking', 'Patient Scheduling', 'PAS Systems', 'Calendar Management',
            'Patient Communication', 'Telephone Skills', 'Data Entry', 'Waiting Lists',
            'NHS Booking', 'Choose & Book', 'e-Referrals', 'Customer Service',
            'Time Management', 'Accuracy', 'Multitasking', 'Microsoft Office'
        ],
        'appointment_admin': [
            'Appointment Administration', 'Scheduling', 'PAS Systems', 'Patient Records',
            'Booking Management', 'Patient Communication', 'Diary Management', 'Data Entry',
            'NHS Administration', 'Waiting Lists', 'Telephone Skills', 'Email Management',
            'Customer Service', 'Attention to Detail', 'Organizational Skills', 'GDPR'
        ],
        'mdt_admin': [
            'MDT Administration', 'Multi-Disciplinary Team', 'Meeting Coordination',
            'Cancer Pathways', 'Clinical Liaison', 'Minute Taking', 'Action Tracking',
            'Patient Lists', 'Outcome Recording', 'Database Management', 'Excel',
            'Communication Skills', 'Time Management', 'Confidentiality', 'Medical Terminology'
        ],
        'medical_secretary': [
            'Medical Secretarial', 'Clinic Correspondence', 'Typing', 'Audio Typing',
            'Appointment Management', 'Patient Records', 'Medical Terminology', 'Letters',
            'Referral Processing', 'Diary Management', 'PAS Systems', 'Microsoft Office',
            'Confidentiality', 'GDPR', 'Communication', 'Organizational Skills'
        ],
        'data_officer': [
            'Data Quality', 'Information Governance', 'NHS Data', 'Validation',
            'Data Entry', 'Database Management', 'Excel', 'Reporting', 'Accuracy',
            'Audit', 'GDPR Compliance', 'Data Protection', 'PAS Systems', 'Analytics',
            'Attention to Detail', 'Problem Solving', 'Communication', 'Time Management'
        ],
        'nhs_clerical': [
            'NHS Administration', 'Clerical Support', 'Data Entry', 'Filing',
            'Patient Records', 'Reception', 'Telephone Skills', 'Email Management',
            'Appointment Booking', 'Microsoft Office', 'Customer Service', 'Communication',
            'Teamwork', 'Confidentiality', 'GDPR', 'Organizational Skills'
        ],
        
        # HEALTHCARE & CARE
        'healthcare_assistant': [
            'Patient Care', 'Personal Care', 'Vital Signs', 'Observations',
            'Safeguarding', 'Dignity', 'Compassionate Care', 'Clinical Support',
            'Medication Administration', 'Infection Control', 'Health & Safety',
            'Communication Skills', 'Teamwork', 'Record Keeping', 'Mobility Support'
        ],
        'care_worker': [
            'Adult Social Care', 'Personal Care', 'Domiciliary Care', 'Service Users',
            'Care Planning', 'Safeguarding', 'Dignity & Respect', 'Independence',
            'Medication Support', 'Daily Living Activities', 'Mental Capacity Act',
            'Person-Centered Care', 'CQC Standards', 'Manual Handling', 'Care Certificate'
        ],
        
        # EDUCATION
        'teaching_assistant': [
            'Classroom Support', 'SEN/SEND', 'Behavior Management', 'Differentiation',
            'Inclusive Education', 'Child Safeguarding', 'Learning Support',
            'Phonics', 'Numeracy', 'Literacy', 'Early Years', 'Primary Education',
            'Secondary Education', 'Individual Education Plans', 'Assessment'
        ],
        
        # CUSTOMER SERVICE
        'customer_service': [
            'Customer Service', 'Client Relations', 'Problem Solving', 'Communication',
            'Complaint Handling', 'Phone Support', 'Email Support', 'CRM Systems',
            'Conflict Resolution', 'Time Management', 'Team Collaboration',
            'Reception', 'Booking Systems', 'Query Resolution', 'Patient Experience'
        ],
        
        # BUSINESS ADMIN
        'business_admin': [
            'Administration', 'Office Management', 'Data Entry', 'Microsoft Office',
            'Excel', 'Word', 'Outlook', 'Filing Systems', 'Record Keeping',
            'Meeting Coordination', 'Diary Management', 'Process Improvement',
            'GDPR', 'Confidentiality', 'Time Management', 'Organizational Skills'
        ],
        
        # TECH ROLES
        'data_analyst': [
            'Data Analysis', 'Excel Advanced', 'SQL', 'Power BI', 'Tableau',
            'Data Visualization', 'Python', 'Statistics', 'Reporting', 'Dashboard',
            'Data Cleansing', 'Data Mining', 'Business Intelligence', 'KPIs', 'Analytics'
        ],
        'data_scientist': [
            'Data Science', 'Machine Learning', 'Python', 'R', 'Statistical Analysis',
            'Predictive Modeling', 'Big Data', 'SQL', 'Data Visualization', 'Algorithms',
            'TensorFlow', 'Pandas', 'NumPy', 'Deep Learning', 'AI'
        ],
        'software_tester': [
            'Software Testing', 'QA', 'Test Cases', 'Manual Testing', 'Automated Testing',
            'Selenium', 'Bug Tracking', 'Jira', 'Regression Testing', 'UAT',
            'Test Plans', 'Agile', 'Quality Assurance', 'Test Automation', 'Defect Management'
        ],
        'business_analyst': [
            'Business Analysis', 'Requirements Gathering', 'Stakeholder Management',
            'Process Mapping', 'Use Cases', 'Business Intelligence', 'SQL', 'Agile',
            'User Stories', 'Gap Analysis', 'Data Analysis', 'Project Management', 'Documentation'
        ],
        'project_manager': [
            'Project Management', 'Agile', 'Scrum', 'Stakeholder Management', 'Risk Management',
            'Budget Management', 'Team Leadership', 'Planning', 'Delivery', 'PMO',
            'PRINCE2', 'Jira', 'MS Project', 'Communication', 'Problem Solving'
        ],
        'it_support': [
            'IT Support', 'Helpdesk', 'Technical Support', 'Troubleshooting', 'Windows',
            'Active Directory', 'Network Support', 'Hardware', 'Software', 'Customer Service',
            'Ticketing Systems', 'Remote Support', 'Installation', 'Configuration', 'Documentation'
        ],
        'web_developer': [
            'Web Development', 'HTML', 'CSS', 'JavaScript', 'React', 'Node.js',
            'Responsive Design', 'Frontend', 'Backend', 'APIs', 'Git', 'jQuery',
            'Bootstrap', 'Web Design', 'Cross-Browser Compatibility'
        ],
        'python_developer': [
            'Python', 'Django', 'Flask', 'APIs', 'Backend Development', 'SQL',
            'Git', 'Object-Oriented Programming', 'Data Structures', 'Algorithms',
            'RESTful APIs', 'Testing', 'Debugging', 'Automation', 'Web Scraping'
        ],
        
        # BUSINESS & PROFESSIONAL
        'digital_marketing': [
            'Digital Marketing', 'SEO', 'SEM', 'Social Media', 'Content Marketing',
            'Google Analytics', 'Facebook Ads', 'Email Marketing', 'PPC', 'Copywriting',
            'Marketing Strategy', 'Brand Management', 'Campaign Management', 'Analytics'
        ],
        'hr_officer': [
            'Human Resources', 'Recruitment', 'Employee Relations', 'HR Policies',
            'Performance Management', 'Onboarding', 'Employment Law', 'GDPR',
            'Payroll', 'Benefits Administration', 'Training & Development', 'HR Systems'
        ],
        'bookkeeper': [
            'Bookkeeping', 'Accounts Payable', 'Accounts Receivable', 'VAT', 'Payroll',
            'Bank Reconciliation', 'Sage', 'Xero', 'QuickBooks', 'Excel', 'Financial Records',
            'Invoicing', 'Expense Management', 'Month-End', 'Year-End'
        ],
        'team_leader': [
            'Team Leadership', 'People Management', 'Performance Management', 'Coaching',
            'Mentoring', 'Target Setting', 'KPIs', 'Communication', 'Conflict Resolution',
            'Decision Making', 'Delegation', 'Motivation', 'Team Development', 'Planning'
        ]
    }
    
    return keywords.get(career_path, [])


def generate_linkedin_profile(cv_data):
    """
    Generate optimized LinkedIn profile text
    """
    
    personal = cv_data['personal_info']
    
    # LinkedIn Headline (120 chars max)
    headline = f"{cv_data['work_history'][0]['title']} | {', '.join(cv_data['skills'][:2])} | T21 Certified"
    if len(headline) > 120:
        headline = f"{cv_data['work_history'][0]['title']} | T21 Certified Professional"
    
    # About Section
    about = personal['summary'] + "\n\n"
    about += "Core Competencies:\n"
    for skill in cv_data['skills'][:8]:
        about += f"• {skill}\n"
    
    about += "\n💡 Open to opportunities in healthcare, education, and administration."
    
    linkedin = {
        'headline': headline,
        'about': about,
        'skills_to_add': cv_data['skills'],
        'tips': [
            "Add a professional profile photo (increases profile views by 14x)",
            "Turn on 'Open to Work' badge for recruiters",
            "Connect with at least 50 people in your industry",
            "Ask colleagues for skill endorsements",
            "Join relevant LinkedIn groups (NHS Jobs, Healthcare UK, etc.)",
            "Post updates once a week to stay visible"
        ]
    }
    
    return linkedin


def get_t21_qualifications():
    """
    Get list of T21 / TQUK qualifications to auto-add
    """
    
    return [
        {
            'title': 'TQUK Endorsed: Proficient Professional Development Learning Course in Understanding RTT and Hospital Administration (PDLC-01-039)',
            'institution': 'T21 Services UK (TQUK Endorsed)',
            'date': 'Year',
            'description': 'RTT pathway management, NHS 18-week standard, hospital administration, PAS systems, data validation'
        },
        {
            'title': 'Certificate of Completion - RTT Pathway Management',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Comprehensive RTT training, clock management, breach prevention, validation'
        },
        {
            'title': 'TQUK Level 2 Certificate in Customer Service',
            'institution': 'T21 Services UK (TQUK Approved)',
            'date': 'Year',
            'description': 'Understanding customer needs, communication, complaint handling'
        },
        {
            'title': 'TQUK Level 2 Certificate in Adult Social Care',
            'institution': 'T21 Services UK (TQUK Approved)',
            'date': 'Year',
            'description': 'Person-centered care, safeguarding, dignity and respect'
        },
        {
            'title': 'TQUK Level 3 Certificate in Supporting Teaching & Learning',
            'institution': 'T21 Services UK (TQUK Approved)',
            'date': 'Year',
            'description': 'Classroom support, SEN assistance, learning activities'
        },
        {
            'title': 'TQUK Functional Skills in English Level 2',
            'institution': 'T21 Services UK (TQUK Approved)',
            'date': 'Year',
            'description': 'Communication, writing, comprehension'
        },
        {
            'title': 'TQUK Functional Skills in Maths Level 2',
            'institution': 'T21 Services UK (TQUK Approved)',
            'date': 'Year',
            'description': 'Numeracy, problem-solving, data analysis'
        },
        {
            'title': 'TQUK Level 2 Certificate in IT User Skills',
            'institution': 'T21 Services UK (TQUK Approved)',
            'date': 'Year',
            'description': 'Microsoft Office, Excel, Word, data management'
        },
        {
            'title': 'TQUK Level 3 Diploma in Adult Care',
            'institution': 'T21 Services UK (TQUK Approved)',
            'date': 'Year',
            'description': 'Advanced care practice, leadership, person-centered support'
        },
        {
            'title': 'TQUK Level 2 Certificate in Principles of Business Administration',
            'institution': 'T21 Services UK (TQUK Approved)',
            'date': 'Year',
            'description': 'Office procedures, communication, record management'
        },
        {
            'title': 'Care Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': '15 care standards including safeguarding, dignity, communication'
        },
        # PROFESSIONAL & TECH COURSES
        {
            'title': 'Project Management Professional Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Project planning, risk management, stakeholder engagement, Agile/Scrum methodologies'
        },
        {
            'title': 'Data Science Professional Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Python, machine learning, statistical analysis, data visualization'
        },
        {
            'title': 'Data Analysis Professional Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Excel advanced, SQL, Power BI, Tableau, data interpretation'
        },
        {
            'title': 'Software Testing Professional Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Manual testing, automated testing, test case design, bug tracking, QA methodologies'
        },
        {
            'title': 'Business Analysis Professional Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Requirements gathering, process mapping, stakeholder analysis, business intelligence'
        },
        {
            'title': 'Cyber Security Fundamentals',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Information security, threat detection, risk assessment, security best practices'
        },
        {
            'title': 'Digital Marketing Professional Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'SEO, social media marketing, content strategy, Google Analytics, email campaigns'
        },
        {
            'title': 'Python Programming Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Python fundamentals, data structures, object-oriented programming, automation'
        },
        {
            'title': 'Web Development Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'HTML, CSS, JavaScript, responsive design, web frameworks'
        },
        {
            'title': 'Cloud Computing (AWS/Azure) Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Cloud infrastructure, deployment, storage solutions, cloud security'
        },
        {
            'title': 'Database Management (SQL) Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'SQL queries, database design, data modeling, performance optimization'
        },
        {
            'title': 'Agile & Scrum Master Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Scrum framework, sprint planning, agile methodologies, team facilitation'
        },
        {
            'title': 'DevOps Fundamentals Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'CI/CD pipelines, automation, containerization, infrastructure as code'
        },
        {
            'title': 'Microsoft Excel Advanced Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Advanced formulas, pivot tables, macros, VBA, data analysis'
        },
        {
            'title': 'Power BI Data Visualization Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Dashboard creation, DAX, data modeling, report automation'
        },
        {
            'title': 'Leadership & Management Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Team leadership, performance management, conflict resolution, strategic planning'
        },
        {
            'title': 'Human Resources Management Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Recruitment, employee relations, HR policies, performance appraisal'
        },
        {
            'title': 'Bookkeeping & Accounting Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Financial records, accounts payable/receivable, VAT, payroll'
        },
        {
            'title': 'Health & Safety in the Workplace Certificate',
            'institution': 'T21 Services UK',
            'date': 'Year',
            'description': 'Risk assessment, COSHH, fire safety, incident reporting'
        }
    ]


# ============================================
# PDF AND WORD EXPORT FUNCTIONS
# ============================================

def export_cv_to_pdf(cv_html, filename="CV.pdf"):
    """
    Convert HTML CV to PDF using reportlab
    Returns BytesIO object for download
    """
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
        from reportlab.lib.colors import HexColor
        from io import BytesIO
        from html.parser import HTMLParser
        
        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Container for PDF elements
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=HexColor('#2c3e50'),
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=HexColor('#3498db'),
            spaceAfter=6,
            spaceBefore=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=10,
            alignment=TA_JUSTIFY,
            spaceAfter=6
        )
        
        # Simple HTML parser to extract text
        class HTMLTextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []
                self.in_h1 = False
                self.in_h2 = False
                self.in_p = False
                
            def handle_starttag(self, tag, attrs):
                if tag == 'h1':
                    self.in_h1 = True
                elif tag == 'h2':
                    self.in_h2 = True
                elif tag == 'p':
                    self.in_p = True
                    
            def handle_endtag(self, tag):
                if tag == 'h1':
                    self.in_h1 = False
                    self.text_parts.append(('h1_end', ''))
                elif tag == 'h2':
                    self.in_h2 = False
                    self.text_parts.append(('h2_end', ''))
                elif tag == 'p':
                    self.in_p = False
                    self.text_parts.append(('p_end', ''))
                    
            def handle_data(self, data):
                data = data.strip()
                if data:
                    if self.in_h1:
                        self.text_parts.append(('h1', data))
                    elif self.in_h2:
                        self.text_parts.append(('h2', data))
                    elif self.in_p:
                        self.text_parts.append(('p', data))
        
        # Parse HTML
        parser = HTMLTextExtractor()
        parser.feed(cv_html)
        
        # Build PDF
        for tag, text in parser.text_parts:
            if tag == 'h1':
                story.append(Paragraph(text, title_style))
            elif tag == 'h2':
                story.append(Paragraph(text, heading_style))
            elif tag == 'p' and text:
                story.append(Paragraph(text, body_style))
            elif tag in ['h1_end', 'h2_end', 'p_end']:
                story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
        
    except ImportError:
        return None  # reportlab not installed


def export_cv_to_word(cv_html, filename="CV.docx"):
    """
    Convert HTML CV to Word document using python-docx
    Returns BytesIO object for download
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from html.parser import HTMLParser
        
        # Create Word document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Simple HTML parser
        class HTMLToWordParser(HTMLParser):
            def __init__(self, doc):
                super().__init__()
                self.doc = doc
                self.in_h1 = False
                self.in_h2 = False
                self.in_p = False
                
            def handle_starttag(self, tag, attrs):
                if tag == 'h1':
                    self.in_h1 = True
                elif tag == 'h2':
                    self.in_h2 = True
                elif tag == 'p':
                    self.in_p = True
                    
            def handle_endtag(self, tag):
                self.in_h1 = False
                self.in_h2 = False
                self.in_p = False
                    
            def handle_data(self, data):
                data = data.strip()
                if data:
                    if self.in_h1:
                        p = self.doc.add_paragraph(data)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.runs[0]
                        run.font.size = Pt(24)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(44, 62, 80)
                    elif self.in_h2:
                        p = self.doc.add_paragraph(data)
                        run = p.runs[0]
                        run.font.size = Pt(14)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(52, 152, 219)
                    elif self.in_p:
                        p = self.doc.add_paragraph(data)
                        run = p.runs[0]
                        run.font.size = Pt(10)
        
        # Parse HTML and build Word doc
        parser = HTMLToWordParser(doc)
        parser.feed(cv_html)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except ImportError:
        return None  # python-docx not installed
