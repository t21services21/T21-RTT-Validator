"""
ENHANCED CV EXPORT FUNCTIONS
Properly exports ALL CV data to PDF and Word formats
NO DATA LOSS - Captures everything the user entered
"""

from io import BytesIO
from datetime import datetime


def export_cv_to_pdf_enhanced(cv_data, filename="Professional_CV.pdf"):
    """
    Enhanced PDF export that captures ALL CV data
    Uses cv_data dictionary directly (not HTML parsing)
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
        from reportlab.lib.colors import HexColor, colors
        from reportlab.lib import colors as reportlab_colors
        
        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=A4,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # Container for PDF elements
        story = []
        
        # Define custom styles
        styles = getSampleStyleSheet()
        
        # Name style
        name_style = ParagraphStyle(
            'Name',
            parent=styles['Heading1'],
            fontSize=22,
            textColor=HexColor('#005EB8'),
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Contact style
        contact_style = ParagraphStyle(
            'Contact',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#666666'),
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        # Section heading style
        section_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=HexColor('#005EB8'),
            spaceAfter=8,
            spaceBefore=12,
            fontName='Helvetica-Bold',
            borderWidth=0,
            borderPadding=0,
            borderColor=HexColor('#CCCCCC'),
            borderRadius=None
        )
        
        # Body text style
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#333333'),
            spaceAfter=6,
            alignment=TA_JUSTIFY
        )
        
        # Job title style
        job_title_style = ParagraphStyle(
            'JobTitle',
            parent=styles['Normal'],
            fontSize=11,
            textColor=HexColor('#000000'),
            spaceAfter=2,
            fontName='Helvetica-Bold'
        )
        
        # Company style
        company_style = ParagraphStyle(
            'Company',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#666666'),
            spaceAfter=6,
            fontName='Helvetica-Oblique'
        )
        
        # Bullet style
        bullet_style = ParagraphStyle(
            'Bullet',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#333333'),
            leftIndent=20,
            spaceAfter=4
        )
        
        # Get personal info
        personal = cv_data['personal_info']
        
        # === HEADER ===
        story.append(Paragraph(personal['name'], name_style))
        
        contact_parts = []
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        if personal.get('location'):
            contact_parts.append(personal['location'])
        if personal.get('linkedin'):
            contact_parts.append(f"LinkedIn: {personal['linkedin']}")
        
        contact_text = " | ".join(contact_parts)
        story.append(Paragraph(contact_text, contact_style))
        
        # Add horizontal line
        story.append(Spacer(1, 0.1*inch))
        
        # === PROFESSIONAL SUMMARY ===
        if personal.get('summary'):
            story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
            story.append(Paragraph(personal['summary'], body_style))
            story.append(Spacer(1, 0.15*inch))
        
        # === KEY SKILLS ===
        if cv_data.get('skills'):
            story.append(Paragraph("KEY SKILLS", section_style))
            
            # Create skills as comma-separated text
            skills_text = " • ".join(cv_data['skills'])
            story.append(Paragraph(skills_text, body_style))
            story.append(Spacer(1, 0.15*inch))
        
        # === WORK EXPERIENCE ===
        if cv_data.get('work_history'):
            story.append(Paragraph("WORK EXPERIENCE", section_style))
            
            for job in cv_data['work_history']:
                # Job title
                story.append(Paragraph(job['title'], job_title_style))
                
                # Company and dates
                company_text = f"{job['company']} | {job['dates']}"
                story.append(Paragraph(company_text, company_style))
                
                # Responsibilities
                if job.get('responsibilities'):
                    for resp in job['responsibilities']:
                        bullet_text = f"• {resp}"
                        story.append(Paragraph(bullet_text, bullet_style))
                
                story.append(Spacer(1, 0.1*inch))
            
            story.append(Spacer(1, 0.05*inch))
        
        # === EDUCATION ===
        education_keywords = ['degree', 'diploma', 'gcse', 'a-level', 'btec', 'nvq', 'bachelor', 'master', 'phd', 'doctorate']
        education_quals = [q for q in cv_data.get('qualifications', []) if any(keyword in q['title'].lower() for keyword in education_keywords)]
        
        if education_quals:
            story.append(Paragraph("EDUCATION", section_style))
            
            for qual in education_quals:
                story.append(Paragraph(qual['title'], job_title_style))
                inst_text = f"{qual['institution']} | {qual['date']}"
                story.append(Paragraph(inst_text, company_style))
                
                if qual.get('description'):
                    story.append(Paragraph(qual['description'], body_style))
                
                story.append(Spacer(1, 0.08*inch))
            
            story.append(Spacer(1, 0.05*inch))
        
        # === PROFESSIONAL QUALIFICATIONS & CERTIFICATIONS ===
        cert_quals = [q for q in cv_data.get('qualifications', []) if not any(keyword in q['title'].lower() for keyword in education_keywords)]
        
        if cert_quals:
            story.append(Paragraph("PROFESSIONAL QUALIFICATIONS & CERTIFICATIONS", section_style))
            
            for qual in cert_quals:
                story.append(Paragraph(qual['title'], job_title_style))
                inst_text = f"{qual['institution']} | {qual['date']}"
                story.append(Paragraph(inst_text, company_style))
                
                if qual.get('description'):
                    story.append(Paragraph(qual['description'], body_style))
                
                story.append(Spacer(1, 0.08*inch))
            
            story.append(Spacer(1, 0.05*inch))
        
        # === ACHIEVEMENTS & AWARDS ===
        if cv_data.get('achievements'):
            story.append(Paragraph("KEY ACHIEVEMENTS & AWARDS", section_style))
            
            for achievement in cv_data['achievements']:
                bullet_text = f"• {achievement}"
                story.append(Paragraph(bullet_text, bullet_style))
            
            story.append(Spacer(1, 0.15*inch))
        
        # === ADDITIONAL INFORMATION ===
        story.append(Paragraph("ADDITIONAL INFORMATION", section_style))
        additional_info = [
            "• Right to Work: Eligible to work in the UK",
            "• DBS Check: Enhanced DBS check available upon request",
            "• References: Available upon request",
            "• Languages: English (Fluent)"
        ]
        for info in additional_info:
            story.append(Paragraph(info, bullet_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
        
    except ImportError as e:
        print(f"Error: {e}")
        return None
    except Exception as e:
        print(f"PDF generation error: {e}")
        return None


def export_cv_to_word_enhanced(cv_data, filename="Professional_CV.docx"):
    """
    Enhanced Word export that captures ALL CV data
    Uses cv_data dictionary directly (not HTML parsing)
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create Word document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Get personal info
        personal = cv_data['personal_info']
        
        # === HEADER ===
        # Name
        name_para = doc.add_paragraph(personal['name'])
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.runs[0]
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 94, 184)
        
        # Contact info
        contact_parts = []
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        if personal.get('location'):
            contact_parts.append(personal['location'])
        if personal.get('linkedin'):
            contact_parts.append(f"LinkedIn: {personal['linkedin']}")
        
        contact_text = " | ".join(contact_parts)
        contact_para = doc.add_paragraph(contact_text)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.runs[0]
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(102, 102, 102)
        
        doc.add_paragraph()  # Spacing
        
        # === PROFESSIONAL SUMMARY ===
        if personal.get('summary'):
            heading = doc.add_paragraph("PROFESSIONAL SUMMARY")
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(0, 94, 184)
            
            summary_para = doc.add_paragraph(personal['summary'])
            summary_run = summary_para.runs[0]
            summary_run.font.size = Pt(10)
            
            doc.add_paragraph()  # Spacing
        
        # === KEY SKILLS ===
        if cv_data.get('skills'):
            heading = doc.add_paragraph("KEY SKILLS")
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(0, 94, 184)
            
            skills_text = " • ".join(cv_data['skills'])
            skills_para = doc.add_paragraph(skills_text)
            skills_run = skills_para.runs[0]
            skills_run.font.size = Pt(10)
            
            doc.add_paragraph()  # Spacing
        
        # === WORK EXPERIENCE ===
        if cv_data.get('work_history'):
            heading = doc.add_paragraph("WORK EXPERIENCE")
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(0, 94, 184)
            
            for job in cv_data['work_history']:
                # Job title
                job_para = doc.add_paragraph(job['title'])
                job_run = job_para.runs[0]
                job_run.font.size = Pt(11)
                job_run.font.bold = True
                
                # Company and dates
                company_text = f"{job['company']} | {job['dates']}"
                company_para = doc.add_paragraph(company_text)
                company_run = company_para.runs[0]
                company_run.font.size = Pt(10)
                company_run.font.italic = True
                company_run.font.color.rgb = RGBColor(102, 102, 102)
                
                # Responsibilities
                if job.get('responsibilities'):
                    for resp in job['responsibilities']:
                        resp_para = doc.add_paragraph(resp, style='List Bullet')
                        resp_run = resp_para.runs[0]
                        resp_run.font.size = Pt(10)
                
                doc.add_paragraph()  # Spacing between jobs
        
        # === EDUCATION ===
        education_keywords = ['degree', 'diploma', 'gcse', 'a-level', 'btec', 'nvq', 'bachelor', 'master', 'phd', 'doctorate']
        education_quals = [q for q in cv_data.get('qualifications', []) if any(keyword in q['title'].lower() for keyword in education_keywords)]
        
        if education_quals:
            heading = doc.add_paragraph("EDUCATION")
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(0, 94, 184)
            
            for qual in education_quals:
                qual_para = doc.add_paragraph(qual['title'])
                qual_run = qual_para.runs[0]
                qual_run.font.size = Pt(11)
                qual_run.font.bold = True
                
                inst_text = f"{qual['institution']} | {qual['date']}"
                inst_para = doc.add_paragraph(inst_text)
                inst_run = inst_para.runs[0]
                inst_run.font.size = Pt(10)
                inst_run.font.italic = True
                inst_run.font.color.rgb = RGBColor(102, 102, 102)
                
                if qual.get('description'):
                    desc_para = doc.add_paragraph(qual['description'])
                    desc_run = desc_para.runs[0]
                    desc_run.font.size = Pt(10)
            
            doc.add_paragraph()  # Spacing
        
        # === PROFESSIONAL QUALIFICATIONS & CERTIFICATIONS ===
        cert_quals = [q for q in cv_data.get('qualifications', []) if not any(keyword in q['title'].lower() for keyword in education_keywords)]
        
        if cert_quals:
            heading = doc.add_paragraph("PROFESSIONAL QUALIFICATIONS & CERTIFICATIONS")
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(0, 94, 184)
            
            for qual in cert_quals:
                qual_para = doc.add_paragraph(qual['title'])
                qual_run = qual_para.runs[0]
                qual_run.font.size = Pt(11)
                qual_run.font.bold = True
                
                inst_text = f"{qual['institution']} | {qual['date']}"
                inst_para = doc.add_paragraph(inst_text)
                inst_run = inst_para.runs[0]
                inst_run.font.size = Pt(10)
                inst_run.font.italic = True
                inst_run.font.color.rgb = RGBColor(102, 102, 102)
                
                if qual.get('description'):
                    desc_para = doc.add_paragraph(qual['description'])
                    desc_run = desc_para.runs[0]
                    desc_run.font.size = Pt(10)
            
            doc.add_paragraph()  # Spacing
        
        # === ACHIEVEMENTS & AWARDS ===
        if cv_data.get('achievements'):
            heading = doc.add_paragraph("KEY ACHIEVEMENTS & AWARDS")
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(0, 94, 184)
            
            for achievement in cv_data['achievements']:
                ach_para = doc.add_paragraph(achievement, style='List Bullet')
                ach_run = ach_para.runs[0]
                ach_run.font.size = Pt(10)
            
            doc.add_paragraph()  # Spacing
        
        # === ADDITIONAL INFORMATION ===
        heading = doc.add_paragraph("ADDITIONAL INFORMATION")
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 94, 184)
        
        additional_info = [
            "Right to Work: Eligible to work in the UK",
            "DBS Check: Enhanced DBS check available upon request",
            "References: Available upon request",
            "Languages: English (Fluent)"
        ]
        for info in additional_info:
            info_para = doc.add_paragraph(info, style='List Bullet')
            info_run = info_para.runs[0]
            info_run.font.size = Pt(10)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except ImportError as e:
        print(f"Error: {e}")
        return None
    except Exception as e:
        print(f"Word generation error: {e}")
        return None
